"""
Student AI Telegram Bot (Google Gemini) — dengan sistem Token/Kredit + ToyyibPay
------------------------------------------------------------------------------
- User baru dapat FREE_STARTING_CREDITS bila /start
- 1 kredit = 1 mesej/soalan
- Bila kredit habis, bot tawar pakej topup (bayar guna ToyyibPay - FPX/kad)
- Admin boleh tambah kredit manual guna /addcredits

Setup:
1. pip install -r requirements.txt
2. Set semua environment variables (lihat README.md / .env.example)
3. python bot.py
"""

import os
import logging
import asyncio
import base64
import io
import re
import urllib.parse

import aiohttp
from google import genai
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ContextTypes,
    filters,
)
from aiohttp import web
from fpdf import FPDF
from docx import Document

import config
import database
import payment
from server import create_web_app

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

TELEGRAM_TOKEN = os.environ["TELEGRAM_BOT_TOKEN"]
GEMINI_API_KEY = os.environ["GEMINI_API_KEY"]

client = genai.Client(api_key=GEMINI_API_KEY)

# Guna beberapa model sebagai fallback — kalau satu kena quota limit / tak available,
# bot automatik cuba model seterusnya. Ni elak bot "mati" bila Google ubah had free-tier
# model tertentu (biasa berlaku, quota free-tier berbeza ikut model & berubah dari semasa
# ke semasa).
GEMINI_MODEL_CANDIDATES = [
    "gemini-flash-lite-latest",  # paling ringan, biasanya quota harian paling tinggi
    "gemini-2.5-flash-lite",
    "gemini-flash-latest",
]


def generate_with_fallback(contents: list[dict]):
    """Cuba setiap model dalam GEMINI_MODEL_CANDIDATES sampai satu berjaya.
    Raise error terakhir kalau semua gagal."""
    last_error = None
    for model_name in GEMINI_MODEL_CANDIDATES:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=contents,
                config={"system_instruction": SYSTEM_PROMPT, "max_output_tokens": 1500},
            )
            return response
        except Exception as e:
            logger.warning(f"Model {model_name} gagal ({e}), cuba model seterusnya...")
            last_error = e
    raise last_error


def clean_markdown(text: str) -> str:
    """Buang markdown asterisk (**bold**) sebab PDF/Word plain text tak render markdown."""
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)


def strip_non_latin1(text: str) -> str:
    """Buang emoji & karakter unicode lain yang tak disokong font PDF asas (Helvetica).
    Word (docx) tak perlu ni sebab dia support unicode penuh."""
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def build_pdf_bytes(title: str, body: str) -> bytes:
    body = strip_non_latin1(clean_markdown(body))
    title = strip_non_latin1(title)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 10, title)
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 12)
    for line in body.split("\n"):
        pdf.set_x(pdf.l_margin)  # reset cursor — elak bug fpdf2 "not enough horizontal space"
        if line.strip() == "":
            pdf.ln(4)
        else:
            pdf.multi_cell(0, 8, line)
    output = pdf.output()
    return bytes(output)


def build_docx_bytes(title: str, body: str) -> bytes:
    body = clean_markdown(body)
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

SYSTEM_PROMPT = """Kau ialah pembantu AI untuk student (StudyBot). Tugas kau:
- Jawab soalan akademik dengan jelas dan tepat, dalam Bahasa Malaysia atau English ikut bahasa student tanya.
- Bila student minta tolong assignment, JANGAN terus bagi jawapan siap. Terangkan konsep, tunjuk langkah-langkah,
  bagi contoh serupa, supaya student faham dan boleh siapkan sendiri. Kalau student betul-betul stuck lepas
  cuba faham, baru bagi jawapan penuh dengan penjelasan.
- Bila diminta ringkaskan nota, buat ringkasan padat dalam bentuk bullet points, senang diingati untuk exam.
- Guna bahasa mesra, ringkas, dan sesuai untuk pelajar (secondary school / university level).
- Kalau soalan tak jelas subjek/tahap apa, tanya sikit untuk clarify sebelum jawab panjang.
"""

user_histories: dict[int, list[dict]] = {}
MAX_HISTORY_MESSAGES = 10
MAX_FILE_SIZE_MB = 20  # had Telegram Bot API (download fail >20MB tak boleh)


# ---------- Helpers ----------

def build_gemini_contents(history: list[dict]) -> list[dict]:
    contents = []
    for msg in history:
        role = "model" if msg["role"] == "assistant" else "user"
        contents.append({"role": role, "parts": [{"text": msg["content"]}]})
    return contents


def topup_keyboard() -> InlineKeyboardMarkup:
    buttons = [
        [
            InlineKeyboardButton(
                f"{p['name']} — {p['credits']} kredit (RM{p['price_myr']:.2f})",
                callback_data=f"topup:{p['id']}",
            )
        ]
        for p in config.PACKAGES
    ]
    return InlineKeyboardMarkup(buttons)


async def process_media(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    file_id: str,
    mime_type: str,
    default_prompt: str,
):
    """Fungsi kongsi untuk handle gambar/PDF/video — download dari Telegram,
    hantar ke Gemini sebagai inline data, tolak 1 kredit macam mesej teks biasa."""
    user = update.effective_user
    user_id = user.id

    await database.ensure_user(user_id, user.username or user.first_name, config.FREE_STARTING_CREDITS)

    if not await database.try_deduct_credit(user_id, amount=1):
        await update.message.reply_text(
            "⚠️ Kredit kau dah habis!\n\nGuna /topup untuk tambah kredit dan sambung belajar 📚",
            reply_markup=topup_keyboard(),
        )
        return

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")

    try:
        tg_file = await context.bot.get_file(file_id)

        if tg_file.file_size and tg_file.file_size > MAX_FILE_SIZE_MB * 1024 * 1024:
            await database.add_credits(user_id, 1)  # refund, bukan salah user
            await update.message.reply_text(
                f"⚠️ Fail ni terlalu besar (max {MAX_FILE_SIZE_MB}MB untuk bot Telegram). "
                "Cuba compress atau hantar bahagian yang lebih kecil ye."
            )
            return

        file_bytes = bytes(await tg_file.download_as_bytearray())
        b64_data = base64.b64encode(file_bytes).decode("utf-8")

        user_caption = (update.message.caption or "").strip() or default_prompt

        contents = [
            {
                "role": "user",
                "parts": [
                    {"inline_data": {"mime_type": mime_type, "data": b64_data}},
                    {"text": user_caption},
                ],
            }
        ]

        response = generate_with_fallback(contents)
        reply_text = response.text or "Maaf, tak dapat proses fail ni. Cuba hantar semula."

    except Exception as e:
        logger.error(f"Gemini media processing error: {e}")
        reply_text = "Alamak, ada masalah nak proses fail kau. Cuba lagi sekejap ye 🙏"
        await database.add_credits(user_id, 1)  # refund

    remaining = await database.get_credits(user_id)
    footer = f"\n\n— 💳 baki kredit: {remaining}"

    for i in range(0, len(reply_text), 4000):
        chunk = reply_text[i : i + 4000]
        is_last = i + 4000 >= len(reply_text)
        await update.message.reply_text(chunk + (footer if is_last else ""))


# ---------- Command Handlers ----------

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    is_new = await database.ensure_user(user.id, user.username or user.first_name, config.FREE_STARTING_CREDITS)
    credits = await database.get_credits(user.id)

    if is_new:
        await update.message.reply_text(
            f"Hai {user.first_name}! Aku StudyBot 🤖📚\n\n"
            f"Kau dapat {config.FREE_STARTING_CREDITS} kredit PERCUMA untuk mula 🎉\n\n"
            "Aku boleh tolong kau:\n"
            "• Jawab soalan pelajaran\n"
            "• Explain & pandu siapkan assignment\n"
            "• Ringkaskan nota panjang\n"
            "• Hantar gambar nota/soalan, PDF, atau video — aku boleh baca sekali! 📷📄🎥\n"
            "• /image <penerangan> — jana gambar AI 🎨\n"
            "• /pdf <soalan> — jawapan terus dalam fail PDF 📄\n"
            "• /doc <soalan> — jawapan terus dalam fail Word 📝\n\n"
            "1 kredit = 1 soalan/mesej. Guna /credits untuk semak baki, "
            "/topup untuk tambah kredit, /reset untuk mula chat baru."
        )
    else:
        await update.message.reply_text(
            f"Hai balik {user.first_name}! Baki kredit kau: {credits} 💳\n"
            "Taip soalan kau, atau /topup kalau nak tambah kredit."
        )


async def credits_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    await database.ensure_user(user_id, update.effective_user.username, config.FREE_STARTING_CREDITS)
    credits = await database.get_credits(user_id)
    await update.message.reply_text(f"💳 Baki kredit kau: {credits}\n\nGuna /topup untuk tambah kredit.")


async def topup_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Pilih pakej kredit yang kau nak beli 👇\n(Bayaran selamat melalui ToyyibPay - FPX/Kad)",
        reply_markup=topup_keyboard(),
    )


async def reset(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_histories.pop(update.effective_user.id, None)
    await update.message.reply_text("Ok, chat history dah clear. Mula baru! 🆕")


async def addcredits_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Admin je boleh guna: /addcredits <user_id> <amount>"""
    caller_id = update.effective_user.id
    if caller_id not in config.ADMIN_USER_IDS:
        await update.message.reply_text("Command ni untuk admin je 🙅")
        return

    args = context.args
    if len(args) != 2 or not args[0].isdigit() or not args[1].lstrip("-").isdigit():
        await update.message.reply_text("Format: /addcredits <user_id> <amount>")
        return

    target_id, amount = int(args[0]), int(args[1])
    await database.add_credits(target_id, amount)
    new_balance = await database.get_credits(target_id)
    await update.message.reply_text(f"Done. User {target_id} baki kredit sekarang: {new_balance}")

    try:
        await context.bot.send_message(
            chat_id=target_id, text=f"🎁 Kau dapat {amount} kredit dari admin! Baki sekarang: {new_balance}"
        )
    except Exception:
        pass  # user mungkin belum start chat dengan bot


# ---------- Generate Content Commands (gambar / PDF / Word) ----------

async def image_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Generate gambar AI guna Pollinations.ai (percuma, model FLUX)."""
    user = update.effective_user
    user_id = user.id
    prompt_text = " ".join(context.args).strip()

    if not prompt_text:
        await update.message.reply_text(
            "Guna format: /image <penerangan gambar>\n"
            "Contoh: /image gambarajah kitaran air untuk kelas sains"
        )
        return

    await database.ensure_user(user_id, user.username or user.first_name, config.FREE_STARTING_CREDITS)

    if not await database.try_deduct_credit(user_id, amount=1):
        await update.message.reply_text(
            "⚠️ Kredit kau dah habis!\n\nGuna /topup untuk tambah kredit dan sambung belajar 📚",
            reply_markup=topup_keyboard(),
        )
        return

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="upload_photo")

    try:
        encoded_prompt = urllib.parse.quote(prompt_text)
        image_url = (
            f"https://image.pollinations.ai/prompt/{encoded_prompt}"
            f"?width=1024&height=1024&nologo=true&model=flux"
        )
        async with aiohttp.ClientSession() as session:
            async with session.get(image_url, timeout=aiohttp.ClientTimeout(total=60)) as resp:
                if resp.status != 200:
                    raise RuntimeError(f"Pollinations return status {resp.status}")
                image_bytes = await resp.read()

        remaining = await database.get_credits(user_id)
        await update.message.reply_photo(
            photo=io.BytesIO(image_bytes),
            caption=f"🎨 \"{prompt_text}\"\n\n— 💳 baki kredit: {remaining}",
        )
    except Exception as e:
        logger.error(f"Pollinations image gen error: {e}")
        await database.add_credits(user_id, 1)  # refund
        await update.message.reply_text("Alamak, gagal jana gambar. Cuba lagi sekejap ye 🙏")


async def generate_answer_text(prompt_text: str) -> str:
    """Panggil Gemini untuk dapatkan jawapan teks (dipakai oleh /pdf dan /doc)."""
    response = generate_with_fallback([{"role": "user", "parts": [{"text": prompt_text}]}])
    return response.text or "Tak dapat jana jawapan. Cuba tanya lain cara."


async def pdf_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Jana jawapan AI terus dalam bentuk fail PDF."""
    user = update.effective_user
    user_id = user.id
    prompt_text = " ".join(context.args).strip()

    if not prompt_text:
        await update.message.reply_text(
            "Guna format: /pdf <soalan atau topik>\nContoh: /pdf ringkaskan bab fotosintesis"
        )
        return

    await database.ensure_user(user_id, user.username or user.first_name, config.FREE_STARTING_CREDITS)

    if not await database.try_deduct_credit(user_id, amount=1):
        await update.message.reply_text(
            "⚠️ Kredit kau dah habis!\n\nGuna /topup untuk tambah kredit dan sambung belajar 📚",
            reply_markup=topup_keyboard(),
        )
        return

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="upload_document")

    try:
        answer_text = await generate_answer_text(prompt_text)
        pdf_bytes = build_pdf_bytes("StudyBot - Jawapan", answer_text)
        remaining = await database.get_credits(user_id)
        await update.message.reply_document(
            document=io.BytesIO(pdf_bytes),
            filename="studybot_jawapan.pdf",
            caption=f"📄 Ni PDF untuk: \"{prompt_text}\"\n\n— 💳 baki kredit: {remaining}",
        )
    except Exception as e:
        logger.error(f"PDF generation error: {e}")
        await database.add_credits(user_id, 1)  # refund
        await update.message.reply_text("Alamak, gagal jana PDF. Cuba lagi sekejap ye 🙏")


async def doc_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Jana jawapan AI terus dalam bentuk fail Word (.docx)."""
    user = update.effective_user
    user_id = user.id
    prompt_text = " ".join(context.args).strip()

    if not prompt_text:
        await update.message.reply_text(
            "Guna format: /doc <soalan atau topik>\nContoh: /doc buatkan nota sejarah kemerdekaan Malaysia"
        )
        return

    await database.ensure_user(user_id, user.username or user.first_name, config.FREE_STARTING_CREDITS)

    if not await database.try_deduct_credit(user_id, amount=1):
        await update.message.reply_text(
            "⚠️ Kredit kau dah habis!\n\nGuna /topup untuk tambah kredit dan sambung belajar 📚",
            reply_markup=topup_keyboard(),
        )
        return

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="upload_document")

    try:
        answer_text = await generate_answer_text(prompt_text)
        docx_bytes = build_docx_bytes("StudyBot - Jawapan", answer_text)
        remaining = await database.get_credits(user_id)
        await update.message.reply_document(
            document=io.BytesIO(docx_bytes),
            filename="studybot_jawapan.docx",
            caption=f"📄 Ni Word doc untuk: \"{prompt_text}\"\n\n— 💳 baki kredit: {remaining}",
        )
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        await database.add_credits(user_id, 1)  # refund
        await update.message.reply_text("Alamak, gagal jana fail Word. Cuba lagi sekejap ye 🙏")


# ---------- Callback Query (topup button) ----------

async def topup_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    package_id = query.data.split(":", 1)[1]
    package = config.get_package(package_id)
    if not package:
        await query.edit_message_text("Pakej tak dijumpai, cuba /topup semula.")
        return

    user_id = query.from_user.id
    await query.edit_message_text(f"Sedang jana link pembayaran untuk {package['name']}...")

    try:
        bill = await payment.create_bill(user_id, package)
        await database.create_transaction(user_id, bill["bill_code"], package)
        await query.message.reply_text(
            f"💳 Pakej: {package['name']} — {package['credits']} kredit (RM{package['price_myr']:.2f})\n\n"
            f"Klik link untuk bayar:\n{bill['url']}\n\n"
            "Kredit akan masuk automatik lepas payment berjaya ✅"
        )
    except Exception as e:
        logger.error(f"Gagal create bill: {e}")
        await query.message.reply_text("Alamak, gagal jana link pembayaran. Cuba lagi sekejap ye 🙏")


# ---------- Media Handlers (gambar / PDF / video) ----------

async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Telegram hantar beberapa saiz — ambil resolusi paling tinggi (last dalam list)
    photo = update.message.photo[-1]
    await process_media(
        update, context, photo.file_id, "image/jpeg",
        default_prompt=(
            "Ni gambar nota/soalan student. Baca dan jawab/terangkan ikut apa yang "
            "sesuai — kalau soalan, pandu step-by-step; kalau nota, boleh ringkaskan "
            "kalau nampak sesuai."
        ),
    )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    doc = update.message.document
    mime_type = doc.mime_type or "application/octet-stream"

    # Gemini paling stabil untuk PDF; jenis lain (docx, pptx, dll) mungkin tak disokong penuh
    if mime_type != "application/pdf":
        await update.message.reply_text(
            "⚠️ Buat masa ni bot ni sokong penuh untuk fail **PDF** je. "
            "Untuk jenis fail lain, cuba export/save jadi PDF dulu ye."
        )
        return

    await process_media(
        update, context, doc.file_id, mime_type,
        default_prompt=(
            "Ni fail PDF nota/bahan belajar student. Ringkaskan isi penting dalam bentuk "
            "bullet points yang senang diingati untuk exam, kecuali student minta benda lain "
            "dalam caption."
        ),
    )


async def handle_video(update: Update, context: ContextTypes.DEFAULT_TYPE):
    video = update.message.video
    mime_type = video.mime_type or "video/mp4"

    await update.message.reply_text("🎥 Video diterima, sedang proses (mungkin ambil sedikit masa lebih lama)...")

    await process_media(
        update, context, video.file_id, mime_type,
        default_prompt=(
            "Ni video (contoh rakaman kelas/lecture) yang student hantar. Terangkan/ringkaskan "
            "isi penting video ni, kecuali student minta benda lain dalam caption."
        ),
    )


# ---------- Main Text Message Handler ----------

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    user_id = user.id
    user_text = update.message.text

    await database.ensure_user(user_id, user.username or user.first_name, config.FREE_STARTING_CREDITS)

    if not await database.try_deduct_credit(user_id, amount=1):
        await update.message.reply_text(
            "⚠️ Kredit kau dah habis!\n\nGuna /topup untuk tambah kredit dan sambung belajar 📚",
            reply_markup=topup_keyboard(),
        )
        return

    history = user_histories.get(user_id, [])
    history.append({"role": "user", "content": user_text})
    history = history[-MAX_HISTORY_MESSAGES:]

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")

    try:
        response = generate_with_fallback(build_gemini_contents(history))
        reply_text = response.text or "Maaf, tak dapat jana jawapan. Cuba tanya lain cara."
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        reply_text = "Alamak, ada masalah nak proses request kau. Cuba lagi sekejap ye 🙏"
        # refund credit sebab request gagal, bukan salah user
        await database.add_credits(user_id, 1)

    history.append({"role": "assistant", "content": reply_text})
    user_histories[user_id] = history[-MAX_HISTORY_MESSAGES:]

    remaining = await database.get_credits(user_id)
    footer = f"\n\n— 💳 baki kredit: {remaining}"

    for i in range(0, len(reply_text), 4000):
        chunk = reply_text[i : i + 4000]
        is_last = i + 4000 >= len(reply_text)
        await update.message.reply_text(chunk + (footer if is_last else ""))


# ---------- Run bot + web server together ----------

async def main():
    await database.init_db()

    application = ApplicationBuilder().token(TELEGRAM_TOKEN).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("credits", credits_cmd))
    application.add_handler(CommandHandler("topup", topup_cmd))
    application.add_handler(CommandHandler("reset", reset))
    application.add_handler(CommandHandler("addcredits", addcredits_cmd))
    application.add_handler(CommandHandler("image", image_cmd))
    application.add_handler(CommandHandler("pdf", pdf_cmd))
    application.add_handler(CommandHandler("doc", doc_cmd))
    application.add_handler(CallbackQueryHandler(topup_callback, pattern=r"^topup:"))
    application.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    application.add_handler(MessageHandler(filters.VIDEO, handle_video))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    web_app = create_web_app(application.bot)
    runner = web.AppRunner(web_app)
    await runner.setup()
    port = int(os.environ.get("PORT", 8080))
    site = web.TCPSite(runner, "0.0.0.0", port)
    await site.start()
    logger.info(f"Web server (ToyyibPay callback) listening on port {port}")

    async with application:
        await application.start()
        await application.updater.start_polling()
        logger.info("Telegram bot polling started...")
        await asyncio.Event().wait()  # run forever


if __name__ == "__main__":
    asyncio.run(main())
